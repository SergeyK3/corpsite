"""One-off DOCX extraction for WP analysis. Output to docs-work/notes/_extraction-wp-docs-002/."""
from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document

BASE = Path(__file__).resolve().parents[1] / "source-materials"
DOCS_WORK = Path(__file__).resolve().parents[1]
OUT = Path(__file__).resolve().parent / "_extraction-wp-docs-002"
IMG_OUT = OUT / "images"

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def style_name(paragraph) -> str | None:
    try:
        return paragraph.style.name if paragraph.style else None
    except Exception:
        return None


def is_heading(paragraph) -> str | None:
    sn = style_name(paragraph) or ""
    if sn.startswith("Heading"):
        return sn
    text = (paragraph.text or "").strip()
    if not text:
        return None
    if len(text) < 120 and paragraph.runs and all(r.bold for r in paragraph.runs):
        return "BoldLine"
    return None


def extract_docx(path: Path) -> dict:
    doc = Document(str(path))
    rel_slug = path.relative_to(BASE).as_posix().replace("/", "__").replace(" ", "_")
    img_dir = IMG_OUT / rel_slug.replace(".docx", "")
    img_dir.mkdir(parents=True, exist_ok=True)

    structure: list[dict] = []
    paragraphs: list[str] = []
    tables: list[dict] = []

    for i, paragraph in enumerate(doc.paragraphs):
        text = (paragraph.text or "").replace("\xa0", " ").strip()
        heading = is_heading(paragraph)
        has_image = bool(paragraph._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"))
        if heading or text or has_image:
            structure.append(
                {
                    "index": i,
                    "style": style_name(paragraph),
                    "heading": heading,
                    "text": text,
                    "has_image": has_image,
                }
            )
        if text:
            paragraphs.append(text)

    for ti, table in enumerate(doc.tables):
        rows = [[cell.text.replace("\xa0", " ").strip() for cell in row.cells] for row in table.rows]
        tables.append({"index": ti, "rows": rows})

    images: list[dict] = []
    image_refs: list[dict] = []
    with zipfile.ZipFile(path, "r") as zf:
        media = [name for name in zf.namelist() if name.startswith("word/media/")]
        for mi, name in enumerate(sorted(media)):
            data = zf.read(name)
            out_path = img_dir / f"{mi:03d}_{Path(name).name}"
            out_path.write_bytes(data)
            images.append(
                {
                    "archive_path": name,
                    "extracted_to": str(out_path.relative_to(DOCS_WORK)).replace("\\", "/"),
                    "size_bytes": len(data),
                    "ext": Path(name).suffix.lower(),
                }
            )

        root = ET.fromstring(zf.read("word/document.xml"))
        for drawing in root.findall(".//w:drawing", NS):
            descr = None
            for node in drawing.iter():
                if node.tag.endswith("docPr"):
                    descr = node.get("descr") or node.get("title") or node.get("name")
            image_refs.append({"description": descr})

    full_text = "\n".join(paragraphs)
    return {
        "path": str(path.relative_to(Path.cwd())).replace("\\", "/"),
        "filename": path.name,
        "folder": path.parent.name,
        "size_bytes": path.stat().st_size,
        "paragraph_count": len(doc.paragraphs),
        "nonempty_paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "image_count": len(images),
        "word_count_approx": len(re.findall(r"\S+", full_text)),
        "structure": structure,
        "tables": tables,
        "images": images,
        "image_refs_in_xml": image_refs,
        "full_text": full_text,
    }


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    IMG_OUT.mkdir(parents=True, exist_ok=True)
    results = [extract_docx(path) for path in sorted(BASE.rglob("*.docx"))]
    (OUT / "extraction-summary.json").write_text(
        json.dumps(results, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    for item in results:
        stem = Path(item["filename"]).stem
        (OUT / f"{stem}.txt").write_text(item["full_text"], encoding="utf-8")
    for item in results:
        print(
            f"{item['path']}: {item['word_count_approx']} words, "
            f"{item['image_count']} images, {item['table_count']} tables"
        )


if __name__ == "__main__":
    main()
